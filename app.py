# app.py — Metrics Report (Patriot-inspired theme, store-timestamped trends, no logo upload)
# Python 3.12 recommended (runtime.txt -> python-3.12)

import os, io, re, glob, time, base64, hashlib
import numpy as np
import pandas as pd
import streamlit as st

# ---------- Optional libs ----------
try:
    import altair as alt
    HAS_ALTAIR = True
except Exception:
    HAS_ALTAIR = False

try:
    import requests
    HAS_REQUESTS = True
except Exception:
    HAS_REQUESTS = False

# ---------- Page config ----------
st.set_page_config(page_title="Metrics Report", layout="wide")

# ---------------- PATRIOT MOBILE–INSPIRED THEME ----------------
PATRIOT_THEME = {
    "font_url": "https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap",
    "font_stack": "Inter, system-ui, -apple-system, Segoe UI, Roboto, Arial, sans-serif",
    "bg": "#FFFFFF",
    "text": "#0F172A",
    "muted": "#475569",
    "primary": "#0B2D52",   # deep navy
    "accent": "#C8102E",    # patriot red
    "card": "#FFFFFF",
    "border": "#E2E8F0",
    "radius": "12px",
    "shadow": "0 8px 24px rgba(2, 6, 23, 0.06)",
}

def apply_theme(t=PATRIOT_THEME):
    if t.get("font_url"):
        st.markdown(f'<link rel="stylesheet" href="{t["font_url"]}">', unsafe_allow_html=True)

    css = f"""
    <style>
      :root {{
        --bg: {t["bg"]};
        --text: {t["text"]};
        --muted: {t["muted"]};
        --primary: {t["primary"]};
        --accent: {t["accent"]};
        --card: {t["card"]};
        --border: {t["border"]};
        --radius: {t["radius"]};
        --shadow: {t["shadow"]};
      }}

      html, body, .stApp {{
        background: var(--bg);
        color: var(--text);
        font-family: {t["font_stack"]};
      }}

      /* Sticky header bar */
      .pm-header {{
        position: sticky; top: 0; z-index: 50;
        background: var(--primary);
        border-bottom: 1px solid var(--border);
        padding: 10px 0;
      }}
      .pm-wrap {{ width: min(1120px, 92vw); margin: 0 auto; display: flex; align-items: center; gap: 16px; }}
      .pm-title {{ margin: 0; padding: 0; color: #fff; font: 700 22px/1.2 Inter, system-ui, -apple-system, Segoe UI, Roboto, Arial, sans-serif; }}

      /* Buttons */
      .stButton > button {{
        background: var(--accent);
        color: #fff;
        border: 1px solid transparent;
        border-radius: var(--radius);
        box-shadow: var(--shadow);
        font-weight: 600;
      }}
      .stButton > button:hover {{ filter: brightness(0.97); }}

      /* Inputs */
      .stTextInput > div > div input,
      .stNumberInput input,
      .stSelectbox > div > div,
      .stDateInput > div > div,
      .stTextArea textarea {{
        border: 1px solid var(--border);
        border-radius: var(--radius);
      }}

      /* Tabs */
      .stTabs [data-baseweb="tab-list"] button[role="tab"] {{
        background: transparent;
        color: var(--muted);
        border-bottom: 2px solid transparent;
      }}
      .stTabs [data-baseweb="tab-list"] button[role="tab"][aria-selected="true"] {{
        color: var(--primary);
        border-color: var(--primary);
      }}

      /* Metrics cards */
      .stMetric {{
        border: 1px solid var(--border);
        padding: .75rem;
        box-shadow: var(--shadow);
        background: var(--card);
        border-radius: var(--radius);
      }}

      /* Dataframes */
      .stDataFrame, .stTable {{
        border-radius: var(--radius);
        box-shadow: var(--shadow);
      }}
      .stDataFrame table {{
        border: 1px solid var(--border);
        border-radius: var(--radius);
      }}

      /* Links */
      a {{ color: var(--primary); }}
      a:hover {{ color: var(--accent); }}

      /* Section headings with red underline accent */
      h2 {{ border-bottom: 2px solid var(--accent); padding-bottom: 4px; }}
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)

def render_header(title="Metrics Report"):
    st.markdown(
        f'''
        <div class="pm-header">
          <div class="pm-wrap">
            <h1 class="pm-title">{title}</h1>
          </div>
        </div>
        ''',
        unsafe_allow_html=True
    )

apply_theme()
render_header("Metrics Report")
# ---------------- END THEME ----------------

# ================== Generic helpers ==================
def norm(s: str) -> str:
    return re.sub(r"[^a-z0-9]+", "", str(s).lower())

def find_col(df: pd.DataFrame, synonyms) -> str | None:
    cols = list(df.columns)
    norm_map = {norm(c): c for c in cols}
    syn_norm = [norm(x) for x in synonyms if str(x).strip()]
    for s in syn_norm:
        if s in norm_map: return norm_map[s]
    for c in cols:
        nc = norm(c)
        for s in syn_norm:
            if s and (s in nc or nc in s): return c
    return None

def idx_or_default(options, value):
    try: return options.index(value) if value in options else 0
    except Exception: return 0

def read_any(uploaded_or_bytes, name_hint: str | None = None):
    """Read CSV or Excel from file-like, path-like or raw bytes."""
    if isinstance(uploaded_or_bytes, (bytes, bytearray)):
        bio = io.BytesIO(uploaded_or_bytes)
        try:
            bio.seek(0); return pd.read_csv(bio)
        except Exception:
            bio.seek(0); return pd.read_excel(bio)
    name = (getattr(uploaded_or_bytes, "name", None) or name_hint or "").lower()
    try:
        if name.endswith(".csv"):  return pd.read_csv(uploaded_or_bytes)
        if name.endswith(".xlsx"): return pd.read_excel(uploaded_or_bytes, engine="openpyxl")
        if name.endswith(".xls"):  return pd.read_excel(uploaded_or_bytes, engine="xlrd")
        try:    return pd.read_csv(uploaded_or_bytes)
        except:
            if hasattr(uploaded_or_bytes, "seek"): uploaded_or_bytes.seek(0)
            return pd.read_excel(uploaded_or_bytes)
    except Exception as e:
        raise RuntimeError(f"Failed to read file: {e}")

def to_percent(series_like):
    s = pd.Series(series_like).astype(str).str.replace('%', '', regex=False)
    vals = pd.to_numeric(s, errors='coerce')
    d = vals.dropna()
    mx = d.max() if not d.empty else None
    if mx is not None and mx <= 1.0:
        vals = vals * 100.0
    return vals

def parse_duration_to_seconds(x) -> float:
    if pd.isna(x): return np.nan
    s = str(x).strip()
    if re.fullmatch(r"^-?\d+(\.\d+)?$", s): return float(s)
    if ":" in s:
        try:
            parts = [float(p) for p in s.split(":")]
            if len(parts) == 3: h, m, s2 = parts; return h*3600 + m*60 + s2
            if len(parts) == 2: m, s2 = parts; return m*60 + s2
        except Exception:
            return np.nan
    m = re.match(r"(\d+(\.\d+)?)", s)
    return float(m.group(1)) if m else np.nan

def format_seconds(secs: float) -> str:
    if pd.isna(secs): return "N/A"
    secs = float(secs)
    h = int(secs // 3600); m = int((secs % 3600) // 60); s = int(secs % 60)
    return f"{h:d}:{m:02d}:{s:02d}" if h > 0 else f"{m:d}:{s:02d}"

def add_time_columns(df: pd.DataFrame, date_col: str) -> pd.DataFrame:
    dts = pd.to_datetime(df[date_col], errors="coerce")
    out = df.copy()
    out["_DT"] = dts
    out["_DATE"] = dts.dt.normalize()
    out["_WEEK_START"] = dts.dt.to_period("W-MON").dt.start_time
    out["_MONTH_START"] = dts.dt.to_period("M").dt.start_time
    return out

def aggregate_by_period_all_skills(df_time: pd.DataFrame,
                                   skill_col: str,
                                   calls_col: str,
                                   aht_sec_col: str,
                                   rate_pct_series: pd.Series | None,
                                   aband_count_col: str | None,
                                   period_col: str) -> pd.DataFrame:
    tmp = df_time.copy()
    tmp["Calls_num"] = pd.to_numeric(tmp[calls_col], errors="coerce").fillna(0.0)
    tmp["AHT_sec"]   = pd.to_numeric(tmp[aht_sec_col], errors="coerce")

    if aband_count_col and aband_count_col in tmp.columns:
        tmp["_AB_CNT"] = pd.to_numeric(tmp[aband_count_col], errors="coerce")
    if rate_pct_series is not None:
        tmp["_AB_RATE"] = pd.to_numeric(rate_pct_series, errors="coerce")  # 0..100

    g = tmp.groupby([skill_col, period_col], dropna=False)
    calls_sum = g["Calls_num"].sum()

    def _wa(group):
        denom = group["Calls_num"].sum()
        return np.nansum(group["AHT_sec"] * group["Calls_num"]) / denom if denom > 0 else np.nan
    aht_w = g.apply(_wa)

    if "_AB_CNT" in tmp.columns:
        ab_cnt = g["_AB_CNT"].sum()
        ab_pct = (ab_cnt / calls_sum.replace(0, np.nan)) * 100.0
    elif "_AB_RATE" in tmp.columns:
        def _weighted_rate(group):
            r = group["_AB_RATE"]
            return np.nansum(r * group["Calls_num"]) / group["Calls_num"].sum() if group["Calls_num"].sum() > 0 else np.nan
        ab_pct = g.apply(_weighted_rate)
    else:
        ab_pct = pd.Series(index=calls_sum.index, dtype=float)

    out = pd.DataFrame({
        "Skill": [i[0] for i in calls_sum.index],
        "period": [i[1] for i in calls_sum.index],
        "Calls": calls_sum.values,
        "AHT_sec": aht_w.values,
        "Abandon %": ab_pct.values
    })
    out["AHT"] = out["AHT_sec"].apply(format_seconds)
    out.sort_values(["Skill", "period"], inplace=True)
    return out

def build_excel_bytes(sheets: dict[str, pd.DataFrame]) -> bytes:
    try:
        import openpyxl
    except Exception:
        raise RuntimeError("openpyxl not installed. Add it to requirements.txt")
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="openpyxl") as writer:
        for sheet_name, df_ in sheets.items():
            df_.to_excel(writer, index=False, sheet_name=(sheet_name[:31] or "Sheet1"))
    bio.seek(0)
    return bio.read()

# ================== Upload stamping helpers ==================
UPLOADS_DIR = "./_uploads"

def _sanitize_filename(name: str) -> str:
    base = os.path.basename(name or "")
    base = re.sub(r"[^A-Za-z0-9_.-]+", "_", base)
    return base or "file.csv"

def _ensure_uploads_dir():
    try:
        os.makedirs(UPLOADS_DIR, exist_ok=True)
        return True
    except Exception:
        return False

def _uploads_manifest_path():
    return os.path.join(UPLOADS_DIR, "uploads_manifest.csv")

def _append_upload_manifest(path: str, role: str, nbytes: int):
    cols = ["path", "role", "bytes", "added_at"]
    row = {
        "path": path, "role": role, "bytes": nbytes,
        "added_at": time.strftime("%Y-%m-%d %H:%M:%S"),
    }
    try:
        if os.path.exists(_uploads_manifest_path()):
            m = pd.read_csv(_uploads_manifest_path())
        else:
            m = pd.DataFrame(columns=cols)
        m = pd.concat([m, pd.DataFrame([row])], ignore_index=True)
        m.to_csv(_uploads_manifest_path(), index=False)
    except Exception:
        pass

def save_uploaded_copy(raw_bytes: bytes, original_name: str, role: str = "main") -> str | None:
    """Save a timestamped copy of any loaded file (manual upload / URL / local) to ./_uploads."""
    if not raw_bytes or not _ensure_uploads_dir():
        return None
    ts = time.strftime("%Y%m%d_%H%M%S")
    safe_name = _sanitize_filename(original_name or f"{role}.csv")
    out_path = os.path.join(UPLOADS_DIR, f"{ts}_{role}_{safe_name}")
    try:
        with open(out_path, "wb") as f:
            f.write(raw_bytes)
        _append_upload_manifest(out_path, role, len(raw_bytes))
        return out_path
    except Exception:
        return None

# ================== Auto-refresh ==================
with st.sidebar:
    st.header("Auto-refresh")
    auto_refresh = st.checkbox("Enable auto-refresh", value=False, key="auto_refresh_chk")
    refresh_secs = st.number_input("Interval (seconds)", 10, 3600, 60, 5, key="auto_refresh_secs")
    if auto_refresh:
        now = time.time()
        last = st.session_state.get("_last_refresh_ts", 0.0)
        if now - last >= refresh_secs:
            st.session_state["_last_refresh_ts"] = now
            (getattr(st, "rerun", None) or st.experimental_rerun)()
    if st.button("🔄 Reload now", key="manual_reload_btn"):
        st.session_state["_last_refresh_ts"] = time.time()
        (getattr(st, "rerun", None) or st.experimental_rerun)()

# ================== Data Store (historical snapshots) ==================
st.sidebar.header("Data Store (historical)")
store_dir = st.sidebar.text_input("Store folder", value="./_store", key="store_dir")
persist_uploads = st.sidebar.checkbox("Persist uploaded/current dataset to store", value=True, key="persist_uploads")
max_merge = st.sidebar.slider("Max files to merge for trends", min_value=5, max_value=500, value=100, step=5, key="max_merge")

def ensure_store():
    try:
        os.makedirs(store_dir, exist_ok=True)
        return True
    except Exception as e:
        st.sidebar.error(f"Cannot create store dir: {e}")
        return False

def manifest_path():
    return os.path.join(store_dir, "manifest.csv")

def load_manifest() -> pd.DataFrame:
    p = manifest_path()
    if os.path.exists(p):
        try:
            return pd.read_csv(p)
        except Exception:
            return pd.DataFrame(columns=["path","bytes_hash","source","rows","cols","added_at"])
    return pd.DataFrame(columns=["path","bytes_hash","source","rows","cols","added_at"])

def save_manifest(dfm: pd.DataFrame):
    try:
        dfm.to_csv(manifest_path(), index=False)
    except Exception as e:
        st.sidebar.warning(f"Manifest save failed: {e}")

def save_snapshot_bytes(b: bytes, source_label: str = "snapshot") -> tuple[bool, str]:
    if not ensure_store(): return False, "Store not available"
    h = hashlib.sha256(b).hexdigest()
    dfm = load_manifest()
    if (dfm["bytes_hash"] == h).any():
        return False, "Duplicate snapshot (already stored)"
    ts = time.strftime("%Y%m%d_%H%M%S")
    fname = f"{ts}_{h[:8]}.csv"
    path = os.path.join(store_dir, fname)
    try:
        with open(path, "wb") as f:
            f.write(b)
        try:
            tmp = pd.read_csv(io.BytesIO(b)); rows, cols = tmp.shape
        except Exception:
            rows, cols = None, None
        new_row = {
            "path": path, "bytes_hash": h, "source": source_label,
            "rows": rows, "cols": cols, "added_at": time.strftime("%Y-%m-%d %H:%M:%S")
        }
        dfm = pd.concat([dfm, pd.DataFrame([new_row])], ignore_index=True)
        save_manifest(dfm)
        return True, path
    except Exception as e:
        return False, f"Save failed: {e}"

def merge_store_csvs(limit: int) -> pd.DataFrame:
    """Merge up to `limit` most recent CSVs, attaching each file's store timestamp to rows."""
    dfm = load_manifest()
    if dfm.empty: return pd.DataFrame()
    try:
        dfm["_key"] = pd.to_datetime(dfm["added_at"], errors="coerce")
        dfm.sort_values(by=["_key"], inplace=True)
    except Exception:
        dfm.sort_values(by=["path"], inplace=True)
    rows = []
    paths = dfm["path"].dropna().tolist()[-limit:]
    addeds = dfm["added_at"].dropna().tolist()[-limit:]
    for p, added in zip(paths, addeds):
        try:
            with open(p, "rb") as f:
                dft = pd.read_csv(f)
            dft["_STORE_ADDED_AT"] = added  # attach the store timestamp
            rows.append(dft)
        except Exception:
            continue
    if not rows: return pd.DataFrame()
    return pd.concat(rows, ignore_index=True)

col_store_a, col_store_b = st.sidebar.columns(2)
with col_store_a:
    if st.button("📄 View store manifest", key="view_manifest_btn"):
        st.session_state["_show_manifest"] = True
with col_store_b:
    if st.button("🧹 Clear store", key="clear_store_btn"):
        if ensure_store():
            try:
                for p in glob.glob(os.path.join(store_dir, "*.csv")):
                    os.remove(p)
                if os.path.exists(manifest_path()):
                    os.remove(manifest_path())
                st.sidebar.success("Store cleared.")
            except Exception as e:
                st.sidebar.error(f"Clear failed: {e}")

if st.session_state.get("_show_manifest"):
    st.subheader("Store Manifest")
    st.dataframe(load_manifest(), use_container_width=True)

# ================== Uploads manifest controls ==================
st.sidebar.header("Uploads (timestamped copies)")
c_up_a, c_up_b = st.sidebar.columns(2)
with c_up_a:
    if st.button("📄 View uploads manifest", key="view_uploads_manifest_btn"):
        st.session_state["_show_uploads_manifest"] = True
with c_up_b:
    if st.button("🧹 Clear uploads folder", key="clear_uploads_btn"):
        try:
            if os.path.isdir(UPLOADS_DIR):
                for p in glob.glob(os.path.join(UPLOADS_DIR, "*")):
                    os.remove(p)
            st.sidebar.success("Uploads folder cleared.")
        except Exception as e:
            st.sidebar.error(f"Clear failed: {e}")

if st.session_state.get("_show_uploads_manifest"):
    st.subheader("Uploads Manifest")
    try:
        up = pd.read_csv(_uploads_manifest_path()) if os.path.exists(_uploads_manifest_path()) else pd.DataFrame()
        st.dataframe(up, use_container_width=True)
    except Exception as e:
        st.info(f"No manifest yet or failed to read: {e}")

# ================== Data sources (Main) ==================
st.sidebar.header("Main Report — Data Source")
source_type = st.sidebar.radio(
    "Choose source", ["Manual upload", "Public CSV URL", "Local folder (latest *.csv)"],
    index=0, key="main_source_radio"
)
main_url    = st.sidebar.text_input("Main CSV URL", os.getenv("MAIN_CSV_URL", ""), key="main_url")
main_folder = st.sidebar.text_input("Main local folder", "./data", key="main_folder")
main_glob   = st.sidebar.text_input("Main filename pattern", "*.csv", key="main_glob")

def try_fetch_csv_url(url: str) -> tuple[pd.DataFrame | None, dict, bytes | None]:
    if not HAS_REQUESTS:
        return None, {"error": "Package 'requests' not installed. Add it to requirements.txt"}, None
    if not url:
        return None, {"error": "No URL provided"}, None
    try:
        r = requests.get(url, timeout=45)
        if r.status_code >= 300:
            return None, {"error": f"{r.status_code}: {r.text[:300]}", "source": url}, None
        df_ = read_any(r.content, name_hint=url)
        return df_, {"source": url, "bytes": len(r.content)}, r.content
    except Exception as e:
        return None, {"error": str(e), "source": url}, None

def load_latest_local_csv(folder: str, pattern: str = "*.csv") -> tuple[pd.DataFrame | None, dict, bytes | None]:
    try:
        paths = sorted(glob.glob(os.path.join(folder, pattern)), key=lambda p: os.path.getmtime(p))
        if not paths: return None, {"error": f"No files matching {pattern} in {folder}"}, None
        latest = paths[-1]
        with open(latest, "rb") as f: b = f.read()
        df_ = read_any(io.BytesIO(b), name_hint=latest)
        return df_, {"source": latest, "mtime": os.path.getmtime(latest)}, b
    except Exception as e:
        return None, {"error": str(e)}, None

df, source_meta, raw_bytes = None, {}, None
try:
    if source_type == "Public CSV URL":
        df, source_meta, raw_bytes = try_fetch_csv_url(main_url)
        if df is None: st.error(f"URL load failed: {source_meta.get('error','')}"); st.stop()
    elif source_type == "Local folder (latest *.csv)":
        df, source_meta, raw_bytes = load_latest_local_csv(main_folder, main_glob)
        if df is None: st.error(f"Local load failed: {source_meta.get('error','')}"); st.stop()
    else:
        uploaded = st.file_uploader("Main report (CSV/XLSX/XLS)", type=["csv", "xlsx", "xls"], key="main_uploader")
        if uploaded is None:
            st.info("Upload the main CSV/Excel file, or choose another source in the sidebar.")
            st.stop()
        try:
            raw_bytes = uploaded.getvalue()
        except Exception:
            raw_bytes = None
        df = read_any(uploaded); source_meta = {"source": "uploaded file", "name": getattr(uploaded, "name", "")}
except Exception as e:
    st.error(f"Failed to load data: {e}"); st.stop()

if df is None or df.empty:
    st.warning("The main report appears to be empty."); st.stop()

# Tag and stamp current df; also create store-based timestamp column for trends
_ING_TS = time.strftime("%Y-%m-%d %H:%M:%S")
df["_INGESTED_AT"] = _ING_TS
df["_STORE_ADDED_AT"] = _ING_TS  # current file uses its ingest time as the trend date
_saved_main = None
try:
    orig_name = source_meta.get("name") or source_meta.get("source") or "main.csv"
    if raw_bytes:
        _saved_main = save_uploaded_copy(raw_bytes, orig_name, role="main")
except Exception:
    _saved_main = None
if _saved_main:
    st.caption(f"Saved uploaded copy: `{_saved_main}`")

# Normalize: "Abandoned (%rec)" -> "Abandon %"
for c in list(df.columns):
    if norm(c) == norm("Abandoned (%rec)"):
        df.rename(columns={c: "Abandon %"}, inplace=True)

st.caption(f"Loaded main report from: **{source_meta.get('source','(unknown)')}**")
st.subheader("Preview — Main Report (first 20 rows)")
st.dataframe(df.head(20), use_container_width=True)

# ================== Column mapping — no date dropdown ==================
st.subheader("Column Mapping — Main Report")
SKILL_SYNS  = ["skill", "skill name", "skill group", "group", "queue", "split", "team", "program", "department", "dept", "category", "line of business", "lob"]
CALLS_SYNS  = ["calls", "total calls", "calls offered", "offered", "inbound calls", "in calls", "total contacts", "contacts", "total interactions", "volume"]
AGENTS_SYNS = ["agents staffed", "agents", "agent count", "staffed agents", "distinct agents", "distinct agent count", "unique agents"]
AHT_SYNS    = ["aht", "average handle time", "avg handle time", "avg handling time", "avg handle", "aht (s)", "aht (sec)", "talk+hold+acw"]
ABAND_CNT_SYNS = ["abandoned count", "abandoned", "abandon count", "aband count", "abandoned calls"]
ABAND_PCT_SYNS = ["abandon %", "abandoned (%rec)", "abandonment rate", "abandon rate"]

skill_guess     = find_col(df, SKILL_SYNS)
calls_guess     = find_col(df, CALLS_SYNS)
agents_guess    = find_col(df, AGENTS_SYNS)
aht_guess       = find_col(df, AHT_SYNS)
aband_cnt_guess = find_col(df, ABAND_CNT_SYNS)
aband_pct_guess = find_col(df, ABAND_PCT_SYNS)

cols = list(df.columns)
skill_col  = st.selectbox("Skill / Group column", cols, index=idx_or_default(cols, skill_guess or cols[0]))
calls_col  = st.selectbox("Calls column",        cols, index=idx_or_default(cols, calls_guess or cols[0]))
agents_col = st.selectbox("Agents Staffed column (per-skill)", cols, index=idx_or_default(cols, agents_guess or cols[0]))
aht_col    = st.selectbox("AHT column", cols, index=idx_or_default(cols, aht_guess or cols[0]))
abandoned_pct_col = st.selectbox("Abandon % column (optional)", ["<none>"] + cols,
                                 index=idx_or_default(["<none>"]+cols, aband_pct_guess if aband_pct_guess else "<none>"))
abandoned_count_col = st.selectbox("Abandoned (count) column (optional, used if % is missing)", ["<none>"] + cols,
                                   index=idx_or_default(["<none>"]+cols, aband_cnt_guess if aband_cnt_guess else "<none>"))

# Fortress → PM Connect rename in df for consistency
df[skill_col] = df[skill_col].astype(str).str.strip()
df.loc[df[skill_col].str.lower() == "fortress", skill_col] = "PM Connect"

# Skills of interest
default_skills = ["B2B Member Success", "B2B Success Activation", "B2B Success Info", "B2B Success Tech Support",
                  "MS Activation", "MS Info", "MS Loyalty", "MS Tech Support", "PM Connect"]
skills_list = st.text_area("Skills of interest (one per line)", value="\n".join(default_skills))
raw_skills = [s.strip() for s in skills_list.splitlines() if s.strip()]
skills_wanted = []
for s in raw_skills:
    if s.lower() == "fortress": s = "PM Connect"
    if s not in skills_wanted: skills_wanted.append(s)

# ================== Second report (Agents & Total Calls) ==================
st.sidebar.header("Second Report (Agents & Total Calls) — Data Source")
second_source_type = st.sidebar.radio(
    "Choose source", ["Manual upload", "Public CSV URL", "Local folder (latest *.csv)"],
    index=0, key="second_source_radio"
)

def try_fetch_csv_url_second(url):
    df_, meta, b = try_fetch_csv_url(url); return df_, meta, b

def load_latest_local_csv_second(folder, pattern):
    df_, meta, b = load_latest_local_csv(folder, pattern); return df_, meta, b

second_df, second_meta, raw2 = None, {}, None
try:
    if second_source_type == "Public CSV URL":
        url2 = st.sidebar.text_input("2nd CSV URL", os.getenv("SECOND_CSV_URL", ""), key="url2")
        if url2:
            second_df, second_meta, raw2 = try_fetch_csv_url_second(url2)
    elif second_source_type == "Local folder (latest *.csv)":
        fold2 = st.sidebar.text_input("2nd local folder", "./data2", key="fold2")
        pat2  = st.sidebar.text_input("2nd filename pattern", "*.csv", key="pat2")
        second_df, second_meta, raw2 = load_latest_local_csv_second(fold2, pat2)
    else:
        uploaded2 = st.file_uploader("Second report (CSV/XLSX/XLS) — overall totals / no skill filter (optional)",
                                     type=["csv", "xlsx", "xls"], key="second_uploader")
        if uploaded2 is not None:
            try:
                raw2 = uploaded2.getvalue()
            except Exception:
                raw2 = None
            second_df = read_any(uploaded2)
            second_meta = {"source": "uploaded file", "name": getattr(uploaded2, "name", "")}
except Exception as e:
    st.warning(f"Second report load failed: {e}")

if second_df is not None and not second_df.empty:
    # Stamp 2nd df and save copy (URL/local/manual)
    second_df["_INGESTED_AT"] = time.strftime("%Y-%m-%d %H:%M:%S")
    second_df["_STORE_ADDED_AT"] = second_df["_INGESTED_AT"]
    if raw2:
        _saved_second = save_uploaded_copy(raw2, second_meta.get("name") or second_meta.get("source") or "second.csv", role="second")
        if _saved_second:
            st.caption(f"Saved uploaded copy (2nd): `{_saved_second}`")
    for c in list(second_df.columns):
        if norm(c) == norm("Abandoned (%rec)"):
            second_df.rename(columns={c: "Abandon %"}, inplace=True)
    st.caption(f"Loaded 2nd report from: **{second_meta.get('source','uploaded file')}**")
    st.dataframe(second_df.head(10), use_container_width=True)

# ================== Core calculations ==================
calls_num  = pd.to_numeric(df[calls_col],  errors="coerce").fillna(0)
agents_num = pd.to_numeric(df[agents_col], errors="coerce").fillna(0)

rates = None
if abandoned_pct_col != "<none>" and abandoned_pct_col in df.columns:
    rates = to_percent(df[abandoned_pct_col])
aband_count_col_final = None
if rates is None and abandoned_count_col != "<none>" and abandoned_count_col in df.columns:
    aband_count_col_final = abandoned_count_col

# --- Totals (defaults from main report) ---
total_calls = int(calls_num.sum())
total_agents = int(agents_num.sum())
calls_label = "Total Calls (from main report)"
agents_label = "Agents Staffed (sum of per-skill)"

# --- Override from SECOND report when available ---
CALLS_SYNS_ALL  = ["calls", "total calls", "calls offered", "offered", "inbound calls", "contacts", "total contacts", "volume"]
if second_df is not None and not second_df.empty:
    AGENTS_SYNS_MINI = ["agents staffed", "agents", "agent count", "distinct", "unique"]
    agents2_guess = find_col(second_df, AGENTS_SYNS_MINI) or next((c for c in second_df.columns if "agent" in c.lower()), None)
    if agents2_guess:
        total_agents = int(pd.to_numeric(second_df[agents2_guess], errors="coerce").fillna(0).sum())
        agents_label = "Agents Staffed (from 2nd report)"
    calls2_guess = find_col(second_df, CALLS_SYNS_ALL) or next((c for c in second_df.columns if "call" in c.lower() or "offered" in c.lower() or "contact" in c.lower()), None)
    if calls2_guess:
        total_calls = int(pd.to_numeric(second_df[calls2_guess], errors="coerce").fillna(0).sum())
        calls_label = "Total Calls (from 2nd report)"

# --- Abandon % total ---
if aband_count_col_final and total_calls > 0:
    aband_num_total = pd.to_numeric(df[aband_count_col_final], errors="coerce").fillna(0).sum()
    total_abandon_pct = (aband_num_total / total_calls) * 100
elif rates is not None and total_calls > 0:
    total_abandon_pct = ((rates.fillna(0) / 100.0) * calls_num).sum() / total_calls * 100
else:
    total_abandon_pct = None

by_skill_core = pd.DataFrame({
    "SKILL": df[skill_col].astype(str),
    "CALLS": calls_num.astype("Int64"),
    "Agents Staffed": agents_num.astype("Int64"),
    "AHT": df[aht_col].astype(str),
})
by_skill_core["Abandon %"] = (rates.round(2).astype(str) + "%") if rates is not None else "N/A"

# ================== Historical dataset selection ==================
st.markdown("---")
st.header("📦 Historical Dataset")
dataset_scope = st.radio(
    "Use which data for TREND charts?",
    ("Current file only", "Merged historical store", "Current + historical"),
    index=1, key="dataset_scope_radio"
)

historical_df = pd.DataFrame()
if dataset_scope != "Current file only":
    historical_df = merge_store_csvs(limit=max_merge)
    if historical_df.empty:
        st.info("No historical files found in the store yet. Save a snapshot to start building history.")
    else:
        st.success(f"Merged historical files: {len(historical_df)} rows")
        st.dataframe(historical_df.head(10), use_container_width=True)

if dataset_scope == "Current file only":
    analysis_df = df.copy()
elif dataset_scope == "Merged historical store":
    analysis_df = historical_df.copy() if not historical_df.empty else df.copy()
else:
    if historical_df.empty:
        analysis_df = df.copy()
    else:
        # UNION of columns so the store date column stays present
        union_cols = sorted(set(df.columns).union(set(historical_df.columns)))
        analysis_df = pd.concat(
            [historical_df.reindex(columns=union_cols), df.reindex(columns=union_cols)],
            ignore_index=True,
        )

# ================== Save snapshot button ==================
col_snap_a, col_snap_b = st.columns([1,1])
with col_snap_a:
    if st.button("💾 Save snapshot to store", key="save_snapshot_btn"):
        if persist_uploads:
            try:
                if (raw_bytes is None) or (raw_bytes[:4] in [b"%PDF", b"PK\x03\x04"]):
                    bytes_to_save = df.to_csv(index=False).encode("utf-8")
                else:
                    bytes_to_save = raw_bytes
            except Exception:
                bytes_to_save = df.to_csv(index=False).encode("utf-8")
            ok, msg = save_snapshot_bytes(bytes_to_save, source_label=source_meta.get("source","snapshot"))
            if ok: st.success(f"Snapshot saved: {msg}")
            else:  st.info(msg)
        else:
            st.info("Enable 'Persist uploaded/current dataset to store' in the sidebar.")
with col_snap_b:
    if st.button("🔁 Reload store manifest", key="reload_manifest_btn"):
        st.session_state["_show_manifest"] = True

# ================== Filled Report (Core) ==================
md = io.StringIO()
def writeln(s=""): md.write(s + "\n")
writeln("## Autofilled Metrics (Core)\n")
writeln(f"### 3. {calls_label}\n**{total_calls}**\n")
writeln(f"### 4. {agents_label}\n**{total_agents}**\n")
writeln("### 6. Abandon %")
writeln(f"**{(str(round(total_abandon_pct, 2)) + '%') if total_abandon_pct is not None else 'N/A'}**\n")
writeln("### 7. AHT (By Group)")
for sk in skills_wanted:
    mask = by_skill_core["SKILL"].str.lower() == sk.lower()
    val = by_skill_core.loc[mask, "AHT"]
    writeln(f"- **{sk}:** {val.iloc[0] if len(val) else 'Not found in this report'}")
writeln("\n### 8. Abandon % (By Group)")
for sk in skills_wanted:
    mask = by_skill_core["SKILL"].str.lower() == sk.lower()
    val = by_skill_core.loc[mask, "Abandon %"]
    writeln(f"- **{sk}:** {val.iloc[0] if len(val) else 'Not found in this report'}")
report_md = md.getvalue()

st.subheader("Filled Report (Core)")
st.markdown(report_md)

st.subheader("Copy & Paste")
tabs = st.tabs(["📋 Report (Markdown)", "📋 KPIs (CSV)", "📋 By-Skill Core Table (CSV)", "📋 By-Skill Core Table (TSV)", "📋 Preview (first 20 rows CSV)"])
with tabs[0]:
    st.code(report_md, language="markdown")
with tabs[1]:
    kpi_df = pd.DataFrame([{
        "Total Calls": total_calls,
        "Agents Staffed": total_agents,
        "Total Abandon %": (round(total_abandon_pct, 2) if total_abandon_pct is not None else None)
    }])
    disp = kpi_df.copy()
    if disp.loc[0, "Total Abandon %"] is not None:
        disp.loc[0, "Total Abandon %"] = f"{disp.loc[0, 'Total Abandon %']:.2f}%"
    st.dataframe(disp, use_container_width=True)
    st.code(kpi_df.to_csv(index=False), language="text")
with tabs[2]:
    st.dataframe(by_skill_core, use_container_width=True)
    st.code(by_skill_core.to_csv(index=False), language="text")
with tabs[3]:
    st.dataframe(by_skill_core, use_container_width=True)
    st.code(by_skill_core.to_csv(index=False, sep="\t"), language="text")
with tabs[4]:
    prev_csv = df.head(20).to_csv(index=False)
    st.dataframe(df.head(20), use_container_width=True)
    st.code(prev_csv, language="text")

st.subheader("Downloads")
st.download_button("⬇️ Download report (Markdown)", data=report_md.encode("utf-8"),
                   file_name="filled_report_core.md", mime="text/markdown")

st.subheader("By-Skill Table — Core Fields")
st.dataframe(by_skill_core, use_container_width=True)

# ================== Skill Trends (using store timestamp) ==================
st.markdown("---")
st.header("📈 Skill Trends — AHT & Abandon % (Daily / Weekly / Monthly)")
st.caption("Trend dates = when each CSV snapshot was saved to the store (or the current file's ingest time).")

trend_df = analysis_df.copy()
if "_STORE_ADDED_AT" not in trend_df.columns:
    trend_df["_STORE_ADDED_AT"] = _ING_TS

# Normalize skills
if skill_col in trend_df.columns:
    trend_df[skill_col] = trend_df[skill_col].astype(str).str.strip()
    trend_df.loc[trend_df[skill_col].str.lower() == "fortress", skill_col] = "PM Connect"

trend_df["_AHT_sec"] = trend_df[aht_col].apply(parse_duration_to_seconds) if aht_col in trend_df.columns else np.nan
trend_df = add_time_columns(trend_df, "_STORE_ADDED_AT")

rate_series = None
if abandoned_pct_col != "<none>" and abandoned_pct_col in trend_df.columns:
    rate_series = to_percent(trend_df[abandoned_pct_col])
aband_count_col_final_for_trend = abandoned_count_col if (abandoned_count_col != "<none>" and abandoned_count_col in trend_df.columns) else None

if skill_col not in trend_df.columns or calls_col not in trend_df.columns:
    st.info("Trends need at least the Skill and Calls columns present in the merged dataset.")
else:
    daily_all   = aggregate_by_period_all_skills(trend_df, skill_col, calls_col, "_AHT_sec", rate_series, aband_count_col_final_for_trend, "_DATE")
    weekly_all  = aggregate_by_period_all_skills(trend_df, skill_col, calls_col, "_AHT_sec", rate_series, aband_count_col_final_for_trend, "_WEEK_START")
    monthly_all = aggregate_by_period_all_skills(trend_df, skill_col, calls_col, "_AHT_sec", rate_series, aband_count_col_final_for_trend, "_MONTH_START")

    all_skills_sorted = sorted([str(x) for x in trend_df[skill_col].dropna().unique()]) if skill_col in trend_df.columns else []
    if not all_skills_sorted:
        st.info("No skill data found to plot. Check your column mapping and historical files.")
    else:
        skill_choice = st.selectbox("Single skill", all_skills_sorted,
                                    index=idx_or_default(all_skills_sorted, "PM Connect" if "PM Connect" in all_skills_sorted else all_skills_sorted[0]),
                                    key="single_skill_select")

        def filter_skill(df_in, skill):
            return df_in[df_in["Skill"].astype(str).str.lower() == skill.lower()].copy()

        daily   = filter_skill(daily_all,   skill_choice)
        weekly  = filter_skill(weekly_all,  skill_choice)
        monthly = filter_skill(monthly_all, skill_choice)

        def delta_str(series):
            if len(series) < 2 or pd.isna(series.iloc[-2]) or pd.isna(series.iloc[-1]): return "—"
            diff = series.iloc[-1] - series.iloc[-2]
            sign = "▲" if diff > 0 else ("▼" if diff < 0 else "—")
            return f"{sign} {diff:.2f}"

        k1, k2, k3, k4 = st.columns(4)
        with k1:
            st.metric("Last AHT (Daily)", daily["AHT"].iloc[-1] if not daily.empty else "N/A",
                      delta=delta_str(daily["AHT_sec"]) if not daily.empty else "—")
        with k2:
            st.metric("Last Abandon % (Daily)", f"{daily['Abandon %'].iloc[-1]:.2f}%" if (not daily.empty and pd.notna(daily['Abandon %'].iloc[-1])) else "N/A",
                      delta=delta_str(daily["Abandon %"]) if not daily.empty else "—")
        with k3:
            st.metric("Last AHT (Weekly)", weekly["AHT"].iloc[-1] if not weekly.empty else "N/A",
                      delta=delta_str(weekly["AHT_sec"]) if not weekly.empty else "—")
        with k4:
            st.metric("Last Abandon % (Weekly)", f"{weekly['Abandon %'].iloc[-1]:.2f}%" if (not weekly.empty and pd.notna(weekly['Abandon %'].iloc[-1])) else "N/A",
                      delta=delta_str(weekly["Abandon %"]) if not weekly.empty else "—")

        # ---------- chart helpers ----------
        def alt_line_chart(df_in: pd.DataFrame, y_col: str, y_title: str):
            df_plot = df_in.copy()
            df_plot[y_col] = pd.to_numeric(df_plot[y_col], errors="coerce")
            chart = (
                alt.Chart(df_plot)
                .mark_line(point=True)
                .encode(
                    x=alt.X("period:T", title="Period"),
                    y=alt.Y(f"{y_col}:Q", title=y_title),
                    tooltip=[
                        alt.Tooltip("period:T", title="Period"),
                        alt.Tooltip("Calls:Q", title="Calls", format=",.0f"),
                        alt.Tooltip("AHT:N", title="AHT"),
                        alt.Tooltip("Abandon %:Q", title="Abandon %", format=".2f")
                    ]
                )
                .properties(height=280, width="container")
            )
            st.altair_chart(chart, use_container_width=True)

        def st_line_chart(df_in: pd.DataFrame, y_col: str, y_title: str):
            if df_in.empty:
                st.info("No data available."); return
            df_plot = df_in.copy()
            df_plot[y_col] = pd.to_numeric(df_plot[y_col], errors="coerce")
            st.write(y_title)
            t = df_plot[["period", y_col]].set_index("period")
            st.line_chart(t)

        def line_chart(df_in: pd.DataFrame, y_col: str, y_title: str):
            if df_in.empty: st.info("No data available."); return
            if HAS_ALTAIR: alt_line_chart(df_in, y_col, y_title)
            else:          st_line_chart(df_in, y_col, y_title)

        # -------- Single-skill charts --------
        st.subheader(f"Daily — {skill_choice}")
        c1, c2 = st.columns(2)
        with c1: line_chart(daily, "AHT_sec", "AHT (seconds)")
        with c2: line_chart(daily, "Abandon %", "Abandon %")
        st.dataframe(daily.assign(**{"Abandon %": daily["Abandon %"].round(2)}), use_container_width=True)

        st.subheader(f"Weekly — {skill_choice}")
        c3, c4 = st.columns(2)
        with c3: line_chart(weekly, "AHT_sec", "AHT (seconds)")
        with c4: line_chart(weekly, "Abandon %", "Abandon %")
        st.dataframe(weekly.assign(**{"Abandon %": weekly["Abandon %"].round(2)}), use_container_width=True)

        st.subheader(f"Monthly — {skill_choice}")
        c5, c6 = st.columns(2)
        with c5: line_chart(monthly, "AHT_sec", "AHT (seconds)")
        with c6: line_chart(monthly, "Abandon %", "Abandon %")
        st.dataframe(monthly.assign(**{"Abandon %": monthly["Abandon %"].round(2)}), use_container_width=True)

        # ---------- Multi-skill compare ----------
        st.markdown("---")
        st.header("🔀 Multi-skill Compare (overlay)")

        all_skills_sorted2 = sorted(daily_all["Skill"].unique().tolist())
        default_preselect = [s for s in ["PM Connect"] if s in all_skills_sorted2] or all_skills_sorted2[:3]
        multi = st.multiselect("Select skills to compare", all_skills_sorted2, default=default_preselect, key="multi_skills_select")

        def overlay_alt(df_in: pd.DataFrame, y_col: str, title: str):
            df_plot = df_in.copy()
            df_plot[y_col] = pd.to_numeric(df_plot[y_col], errors="coerce")
            chart = (
                alt.Chart(df_plot)
                .mark_line(point=True)
                .encode(
                    x=alt.X("period:T", title="Period"),
                    y=alt.Y(f"{y_col}:Q", title=title),
                    color=alt.Color("Skill:N", legend=alt.Legend(title="Skill")),
                    tooltip=[
                        alt.Tooltip("Skill:N"),
                        alt.Tooltip("period:T", title="Period"),
                        alt.Tooltip("Calls:Q", format=",.0f"),
                        alt.Tooltip("AHT:N"),
                        alt.Tooltip("Abandon %:Q", format=".2f")
                    ]
                )
                .properties(height=320, width="container")
            )
            st.altair_chart(chart, use_container_width=True)

        def overlay_st(df_in: pd.DataFrame, y_col: str, title: str):
            st.write(title)
            if df_in.empty: st.info("No data for the selected skills."); return
            df_plot = df_in.copy()
            df_plot[y_col] = pd.to_numeric(df_plot[y_col], errors="coerce")
            p = df_plot.pivot_table(index="period", columns="Skill", values=y_col, aggfunc="mean")
            st.line_chart(p)

        def overlay_chart(df_in: pd.DataFrame, y_col: str, title: str):
            if df_in.empty: st.info("No data for the selected skills."); return
            if HAS_ALTAIR: overlay_alt(df_in, y_col, title)
            else:          overlay_st(df_in, y_col, title)

        if multi:
            d_daily   = daily_all[daily_all["Skill"].isin(multi)].copy()
            d_weekly  = weekly_all[weekly_all["Skill"].isin(multi)].copy()
            d_monthly = monthly_all[daily_all["Skill"].isin(multi)].copy() if False else monthly_all[monthly_all["Skill"].isin(multi)].copy()

            st.subheader("Daily compare")
            oc1, oc2 = st.columns(2)
            with oc1: overlay_chart(d_daily, "AHT_sec", "AHT (seconds)")
            with oc2: overlay_chart(d_daily, "Abandon %", "Abandon %")
            st.dataframe(d_daily.assign(**{"Abandon %": d_daily["Abandon %"].round(2)}), use_container_width=True)

            st.subheader("Weekly compare")
            oc3, oc4 = st.columns(2)
            with oc3: overlay_chart(d_weekly, "AHT_sec", "AHT (seconds)")
            with oc4: overlay_chart(d_weekly, "Abandon %", "Abandon %")
            st.dataframe(d_weekly.assign(**{"Abandon %": d_weekly["Abandon %"].round(2)}), use_container_width=True)

            st.subheader("Monthly compare")
            oc5, oc6 = st.columns(2)
            with oc5: overlay_chart(d_monthly, "AHT_sec", "AHT (seconds)")
            with oc6: overlay_chart(d_monthly, "Abandon %", "Abandon %")
            st.dataframe(d_monthly.assign(**{"Abandon %": d_monthly["Abandon %"].round(2)}), use_container_width=True)

# ================== Word/PDF exports ==================
try:
    from docx import Document
    def build_docx(md_text):
        doc = Document(); doc.add_heading("Autofilled Metrics (Core)", level=1)
        for line in md_text.splitlines():
            if line.startswith("### "): doc.add_heading(line.replace("### ", ""), level=2)
            elif line.startswith("## "): continue
            else:
                if line.strip(): doc.add_paragraph(line)
        bio = io.BytesIO(); doc.save(bio); bio.seek(0); return bio.getvalue()
    st.download_button("⬇️ Download core report (Word .docx)", data=build_docx(report_md),
                       file_name="filled_report_core.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
except Exception as e:
    st.info(f"Word export disabled: {e}")

# Preferred PDF via ReportLab; fallback to fpdf2 if installed
pdf_ready = False
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import inch
    def build_pdf(md_text):
        bio = io.BytesIO(); c = canvas.Canvas(bio, pagesize=letter)
        width, height = letter; L = 0.75 * inch; T = 0.75 * inch; B = 0.75 * inch
        y = height - T; c.setFont("Times-Bold", 14); c.drawString(L, y, "Autofilled Metrics (Core)"); y -= 0.3*inch
        c.setFont("Times-Roman", 11); import textwrap as _tw
        for line in md_text.splitlines():
            if not line.strip(): y -= 0.18*inch
            else:
                for w in _tw.wrap(line, width=95): c.drawString(L, y, w); y -= 0.18*inch
            if y < B: c.showPage(); y = height - T; c.setFont("Times-Roman", 11)
        c.showPage(); c.save(); bio.seek(0); return bio.getvalue()
    pdf_ready = True
except Exception:
    try:
        from fpdf import FPDF
        def build_pdf(md_text):
            pdf = FPDF(); pdf.set_auto_page_break(auto=True, margin=15); pdf.add_page()
            pdf.set_font("Helvetica", size=11)
            for line in md_text.splitlines():
                pdf.multi_cell(0, 6, line if line.strip() else " ")
            return pdf.output(dest="S").encode("latin1")
        pdf_ready = True
    except Exception as e:
        st.info(f"PDF export disabled: {e}")

if pdf_ready:
    st.download_button("⬇️ Download core report (PDF)", data=build_pdf(report_md),
                       file_name="filled_report_core.pdf", mime="application/pdf")
